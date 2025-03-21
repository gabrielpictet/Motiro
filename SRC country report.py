import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

## Motiro Regional report template: prepares the text, plots and graphs for a regional report
## A regional report is the result of a Motiro process in a region, where several teams have conducted the Motiro survey
## and team discussions. When team discussions are completed, the regional report is prepared to summarize and triangulate
## the findings from both the survey and the team discussions. The report is used to identify common patterns and differences.
## This template can be used for circles, branches, or any other structure that has more than one team.

## Enter the name of the structure, i.e. "region", "district", "branch", etc. in the variable below
NS = 'SRC'
NS_name = 'Swiss Red Cross'
country_name = "Switzerland"
structure = "Country"
#structure_name = "Donetsk Region"
structure_name = country_name
region=structure_name.replace(" ", "_")
month = "March"
year = 2025

## Load the data
#df = pd.read_csv(f'{NS}_ALL.csv')
df = pd.read_csv('Individual.csv')

#filter to keep only the data from the region
#df = df[df['region'] == structure_name]
df = df[df['Country'] == structure_name]

# Convert 'Survey Data' column to datetime
df['Survey Data'] = pd.to_datetime(df['Survey Data'])
# Create a new column 'Date' containing only the date portion
df['Date'] = df['Survey Data'].dt.date
# Step 3: Ensure 'Date' column is in desirable datetime format (if needed)
df['Date'] = pd.to_datetime(df['Date']).dt.strftime('%d %B %Y')
# Step 4: Find the most recent date
most_recent_date = df['Date'].max()
earliest_date = df['Date'].min()

# variable for the number of teams surveyed
number_teams_surveyed = df['Team Name'].nunique()

# Create a Word Document
doc = Document()

# Add documnt title (levels 1 and 2)
doc.add_heading(f'Motiro Report: {structure_name}', level=1)
doc.add_heading(f'{NS_name}, {month} {year}', level=2)

doc.add_heading('Introduction', level=2)
doc.add_paragraph(f'The commencement of the full-scale Russian-Ukrainian war in 2022 dramatically changed the context of the humanitarian action in Ukraine as a whole and for the Ukrainian Red Cross Society (URCS) in particular. With humanitarian needs on a drastic rise, the need for volunteers and staff for all the humanitarian actors in the field, including URCS, soared correspondingly. Under the dire circumstances of the overall chaos, mass migration, full-fledged destruction, even retention of the existing volunteers turned into an enormous challenge, let alone recruitment of the new ones. The resulting downturn in the number of volunteers, combined with the aforementioned upsurge of the demand, led URCS to the introduction of the volunteer allowance scheme (already tried out in Ukraine during the COVID-19 pandemic) - id est, paying financial compensation to volunteers based on the number of the hours spent volunteering by a given person (the threshold having been set at 40 hours per month). It was obvious even at the outset, however, that utilizing such a scheme couldn’t provide but a temporary respite. The lack of the necessary funds caused the abandonment of the volunteer allowance scheme by the end of 2024. This development, in turn, raised the issue of thoroughly researching factors influencing motivation of the URCS volunteers, to the purpose of developing the new approach to their motivation - a decidedly non-financial one. This research, therefore, deals with the ways of effective non-financial support and recognition of the URCS volunteers under the changed conditions.')
doc.add_heading(f'Background - State of volunteer and staff Motivation in {structure_name}', level=2)    
doc.add_paragraph('We now present the problems in volunteering and staff motivation, engagement and retention in your region in early 2024 (i.e. even before the Motiro surveys and discussions)? Any peculiar local context and developments relevant to the report? (e.g., closeness to the frontline, aged volunteers, high turnover, disinterested management, etc.)') 

doc.add_heading('Purpose', level=2)
doc.add_paragraph('Nationally, the main objective of the research concerns development and implementation of the 2026-2030 URCS Volunteering Development Strategy amid growing needs in volunteers, on the one hand, and abandonment of the financial incentives, on the other. Yet another objective is identifying the most pressing issues facing URCS volunteers and staff in terms of their well-being; resolution of these should positively impact retention of the personnel. The research aims to extract new valuable insights as to different branch teams, their peculiarities as well as common patterns motivation-wise, analyzing interconnections between different motivational factors and motivational outcomes and identifying already existing good practices regarding personnel of the URCS branches. This, in turn, will enable URCS to create a comprehensive, evidence-based plan of actions to improve motivation, engagement and retention of its staff and volunteers.')

doc.add_paragraph('Any specific things you wanted to learn about in your regions, problems to investigate, solutions to find, challenges to address? If so, write them down.')

doc.add_heading('Method', level=2)
doc.add_heading('Self-determination theory', level=3)
doc.add_paragraph('The Motiro approach and the Motiro app are based on Self-determination theory (SDT) formulated by Ryan and Deci')

# Add SDT inforgraphic picture
doc.add_picture('Infographic SDT.png', width=Inches(6), height=Inches(3))

# Add Text
doc.add_paragraph('The Motiro survey measures several aspects of a successful volunteer experience. All questions – except for the demographic ones – belong to one of the following three conceptual blocks:')
doc.add_paragraph('Outcome variables: the indicators of a successful and satisfying volunteer experience', style='ListBullet')
doc.add_paragraph('Satisfaction of basic needs for autonomy, belongingness, and competence', style='ListBullet')
doc.add_paragraph('Leadership and management practices', style='ListBullet')
doc.add_heading('Outcome variables', level=4)
doc.add_paragraph('We focus on the following six aspects of successful volunteering:')
doc.add_paragraph('Overall satisfaction with the volunteer activity', style='ListBullet')
doc.add_paragraph('Identification with (emotional commitment to) the National Society', style='ListBullet')
doc.add_paragraph('Intent to continue volunteering for the National Society', style='ListBullet')
doc.add_paragraph('Willingness to share one’s ideas for improvement', style='ListBullet')
doc.add_paragraph('Value congruence (the match between volunteers’ values and the values of the National Society)', style='ListBullet')
doc.add_paragraph('Psychological health (e.g. vigor or burnout)', style='ListBullet')
doc.add_paragraph('High levels on these outcome variables are the target to reach with effective leadership and management practices.')
doc.add_heading('Basic psychological needs', level=4)
doc.add_paragraph('To shed light on the volunteer experience, we make use of the self-determination theory. This approach to human motivation explains why some aspects of the volunteer activity and the organization either boost or undermine volunteers’ motivation.')
doc.add_paragraph('Self-determination theory assumes that there are three basic psychological needs for humans. These needs have to be satisfied so that people can thrive in whatever activity they are involved in, such as work, education, health-related behavior, or volunteering:')
doc.add_paragraph('The need for autonomy refers to the desire to feel a sense of psychological freedom and choice in an activity. People want to have a say in what they do and be allowed to voice their opinion.', style='ListBullet')
doc.add_paragraph('The need for belongingness refers to the desire to develop meaningful and warm relationships with other individuals and to feel as part of a group.', style='ListBullet')
doc.add_paragraph('The need for competence addresses the desire to be able to handle an optimally challenging task successfully, and to attain valuable outcomes.', style='ListBullet')
doc.add_paragraph('We use the three needs as early indicators of how fulfilling the volunteer experience is. Prior studies demonstrated that satisfaction versus frustration of these three basic needs explains the effect of leadership and management on work outcomes in the contexts of both paid and voluntary work, including the work for the Red Cross.')
doc.add_heading('Leadership and management practices', level=4)
doc.add_paragraph('The basic psychological needs represent the pathways to a successful and fulfilling volunteer experience. Many features of one’s volunteer activity may either satisfy or thwart the desire for autonomy, belongingness, and competence.')
doc.add_paragraph('The Motiro approach focuses on the impact of leadership and management practices, for example:')
doc.add_paragraph('How much do the supervisors support and encourage volunteers’ autonomy?', style='ListBullet')
doc.add_paragraph('How strongly do volunteers support each other?', style='ListBullet')
doc.add_paragraph('How much do the volunteers feel appreciated for their efforts?', style='ListBullet')
doc.add_paragraph('With respect to the third question, the survey looks at different facets of appreciation, that is, not only rewards, but also feedback about the impact of volunteering. Furthermore, we also address appreciation from family, friends, and the community/neighborhood.')
doc.add_paragraph('Motiro data and findings from the Red Cross Red Crescent strongly suggest that staff motivation also affects the motivation of volunteers, which is why the URCS included staff in the Motiro process.')

# add text - section 2 Methods
doc.add_heading(f'Motiro in {country_name} and in the {structure_name}', level=3)
doc.add_heading(f'The Motiro process in {country_name}', level=4)
doc.add_heading(f'The Motiro process in {structure_name}', level=4)
doc.add_paragraph(f'The Motiro process in {structure_name} consists of two main parts:')
doc.add_paragraph('The Motiro survey', style='ListBullet')
doc.add_paragraph('The team discussions around the survey results', style='ListBullet')

# add text - section 2 data and data limitations
doc.add_heading(f'Data and data limitations observed in {structure_name}', level=2)
doc.add_paragraph('In this section we describe the problems encountered when implementing the Motiro survey and team discussions.')
doc.add_paragraph('Some of these problems may have affected the results.')
doc.add_paragraph('We consider two types of data:')
doc.add_paragraph('The survey data', style='ListBullet')
doc.add_paragraph('The team discussion findings', style='ListBullet')

doc.add_heading(f'Issues and obstacles encountered in {structure_name}', level=3)
doc.add_paragraph('We discuss the limitations of the (1) survey and (2) the team discussions you know of and how you think these limitations influenced the results at team and regional levels.')
doc.add_heading('Issues conducting the survey', level=4)

doc.add_paragraph(f'We conducted a Motiro survey in {number_teams_surveyed} teams between {earliest_date} and {most_recent_date}, the Motiro app has been used by:')
doc.add_paragraph('The survey was conducted in the following teams:')
doc.add_paragraph('Team survey response rates')

doc.add_heading('Potential biases in the data', level=3)
doc.add_heading('Survey Coverage bias', level=4)
doc.add_heading('Survey Self-Selection bias', level=4)
doc.add_heading('Team discussions', level=4)
doc.add_heading('Team discussions limitations', level=3)
doc.add_heading('Team discussion coverage', level=4)
doc.add_heading('Team member self-selection bias', level=4)
doc.add_heading('Discussion biases', level=4)

doc.add_heading('Results', level=2)

doc.add_picture(f'{NS}_{structure_name}_spider.png', width=Inches(6), height=Inches(6))

# General data and illustrations on the region as a whole (analysis of Motiro app graphs and scores)
doc.add_heading(f'Motivational outcomes in {structure_name}', level=3)
doc.add_heading(f'Aggregated survey results in {structure_name}', level=4)

doc.add_paragraph('<Paste from the app the regional motivation dashboard here>')

doc.add_heading(f'Aggregated survey results in {structure_name} teams', level=4)

doc.add_heading(f'Team discussion findings on motivational outcomes', level=4)
doc.add_paragraph(f'We now describe and explain patterns, commonalities and differences between teams in {structure_name}, including important elements of their plans of action')

doc.add_heading(f'Well-being', level=3)
doc.add_heading(f'Survey results', level=4)

doc.add_picture(f'{NS}_{structure_name}_wellbeing_bar.png', width=Inches(6), height=Inches(2.5))

# Load the CSV file into a DataFrame
wellbeing_df = pd.read_csv(f'{NS}_{structure_name}_Well-being.csv')

# Add a table to the document
table = doc.add_table(rows=wellbeing_df.shape[0] + 1, cols=wellbeing_df.shape[1])

# Add the header row
for j, col in enumerate(wellbeing_df.columns):
    cell = table.cell(0, j)
    cell.text = col

    # Set font size to 9 for the header row
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(9)

# Add the data rows
for i in range(wellbeing_df.shape[0]):
    for j in range(wellbeing_df.shape[1]):
        cell = table.cell(i + 1, j)
        cell.text = str(wellbeing_df.iat[i, j])
        
        # Center-align the text if the column contains numbers
        if pd.api.types.is_numeric_dtype(wellbeing_df.iloc[:, j]):
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
      
        # Set font size to 9
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

# Apply table style
table.style = 'Table Grid'

doc.add_heading(f'Team discussion findings', level=4)
doc.add_paragraph(f'Describe and explain patterns, commonalities and differences between teams in {structure_name} including in their plans of action')

doc.add_heading(f'Engagement', level=3) 
doc.add_heading(f'Survey results', level=4)

doc.add_picture(f'{NS}_{structure_name}_engagement_spider.png', width=Inches(3), height=Inches(3))

doc.add_paragraph(f'<Paste the regional bar graph here>')

# Load the CSV file into a DataFrame
engagement_df = pd.read_csv(f'{NS}_{structure_name}_Engagement.csv')

# Add a table to the document
table = doc.add_table(rows=engagement_df.shape[0] + 1, cols=engagement_df.shape[1])

# Add the header row
for j, col in enumerate(engagement_df.columns):
    cell = table.cell(0, j)
    cell.text = col

    # Set font size to 9 for the header row
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(9)

# Add the data rows
for i in range(engagement_df.shape[0]):
    for j in range(engagement_df.shape[1]):
        cell = table.cell(i + 1, j)
        cell.text = str(engagement_df.iat[i, j])
        
        # Center-align the text if the column contains numbers
        if pd.api.types.is_numeric_dtype(engagement_df.iloc[:, j]):
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
       
        # Set font size to 9
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

# Apply table style
table.style = 'Table Grid'

doc.add_heading(f'Team discussion findings in motivational outcomes', level=4)
doc.add_paragraph(f'Describe and explain patterns, commonalities and differences between teams in {structure_name} including in their plans of action')

doc.add_heading(f'The three basic psychological needs and intrinsic motivation', level=3)
doc.add_heading(f'Survey results', level=4)
doc.add_paragraph(f'<Paste the regional intrinsic motivation dashboard here>')

doc.add_picture(f'{NS}_{structure_name}_needs_spider.png', width=Inches(2), height=Inches(2))

# Load the CSV file into a DataFrame
needs_df = pd.read_csv(f'{NS}_{structure_name}_needs.csv')

# Add a table to the document
table = doc.add_table(rows=needs_df.shape[0] + 1, cols=needs_df.shape[1])

# Add the header row
for j, col in enumerate(needs_df.columns):
    cell = table.cell(0, j)
    cell.text = col

    # Set font size to 9 for the header row
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(9)

# Add the data rows
for i in range(needs_df.shape[0]):
    for j in range(needs_df.shape[1]):
        cell = table.cell(i + 1, j)
        cell.text = str(needs_df.iat[i, j])
        
        # Center-align the text if the column contains numbers
        if pd.api.types.is_numeric_dtype(needs_df.iloc[:, j]):
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
       
        # Set font size to 9
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

# Apply table style
table.style = 'Table Grid'

doc.add_heading(f'Autonomy', level=4)

doc.add_picture(f'{NS}_{structure_name}_autonomy_bar.png', width=Inches(6), height=Inches(2.5))

doc.add_heading(f'Key result', level=4)
doc.add_paragraph(f'We now present the lessons and evidence on which we will base our discussion and recommendations.')

doc.add_heading(f'Belonging', level=4) 
doc.add_paragraph(f'<Paste the regional bar graph from Motiro report here>')
doc.add_picture(f'{NS}_{structure_name}_belonging_bar.png', width=Inches(6), height=Inches(2.5))
doc.add_heading(f'Key result', level=4)

doc.add_heading(f'Competence', level=4)
doc.add_paragraph(f'<Paste the regional bar graph from Motiro report here>')
doc.add_picture(f'{NS}_{structure_name}_competence_bar.png', width=Inches(6), height=Inches(2.5))
doc.add_heading(f'Key result', level=4)

doc.add_heading(f'Team discussion findings on intrinsic motivators', level=4)
doc.add_paragraph(f'We now describe and explain patterns, commonalities and differences between teams in the region including relevant elements from their plans of action')

doc.add_heading(f'Leadership', level=3)
doc.add_heading(f'Survey results', level=4)

doc.add_picture(f'{NS}_{structure_name}_leadership_spider.png', width=Inches(3), height=Inches(3))
doc.add_picture(f'{NS}_{structure_name}_leadership_bar.png', width=Inches(6), height=Inches(2.5))
# Load the CSV file into a DataFrame
leadership_df = pd.read_csv(f'{NS}_{structure_name}_leadership.csv')

# Add a table to the document
table = doc.add_table(rows=leadership_df.shape[0] + 1, cols=leadership_df.shape[1])

# Add the header row
for j, col in enumerate(leadership_df.columns):
    cell = table.cell(0, j)
    cell.text = col

    # Set font size to 9 for the header row
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(9)

# Add the data rows
for i in range(leadership_df.shape[0]):
    for j in range(leadership_df.shape[1]):
        cell = table.cell(i + 1, j)
        cell.text = str(leadership_df.iat[i, j])
        
        # Center-align the text if the column contains numbers
        if pd.api.types.is_numeric_dtype(leadership_df.iloc[:, j]):
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
       
        # Set font size to 9
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

# Apply table style
table.style = 'Table Grid'

doc.add_heading(f'Listening', level=4)
doc.add_paragraph(f'<Paste the regional bar graph from Motiro report here>')
doc.add_heading(f'Key result', level=4)

doc.add_heading(f'Understanding', level=4)
doc.add_paragraph(f'<Paste the regional bar graph from Motiro report here>')
doc.add_heading(f'Key result', level=4)

doc.add_heading(f'Encouragement', level=4)
doc.add_paragraph(f'<Paste the regional bar graph from Motiro report here>')
doc.add_heading(f'Key result', level=4)

doc.add_heading(f'Team discussion findings on leadership and leadership skills', level=4)

doc.add_heading(f'Management and extrinsic motivation', level=3)
doc.add_heading(f'Survey results', level=4)

doc.add_picture(f'{NS}_{structure_name}_management_spider.png', width=Inches(3), height=Inches(3))

# Load the CSV file into a DataFrame
management_df = pd.read_csv(f'{NS}_{structure_name}_Management.csv')

# Add a table to the document
table = doc.add_table(rows=management_df.shape[0] + 1, cols=management_df.shape[1])

# Add the header row
for j, col in enumerate(management_df.columns):
    cell = table.cell(0, j)
    cell.text = col

    # Set font size to 9 for the header row
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(9)

# Add the data rows
for i in range(management_df.shape[0]):
    for j in range(management_df.shape[1]):
        cell = table.cell(i + 1, j)
        cell.text = str(management_df.iat[i, j])
        
        # Center-align the text if the column contains numbers
        if pd.api.types.is_numeric_dtype(management_df.iloc[:, j]):
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
      
        # Set font size to 9
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

# Apply table style
table.style = 'Table Grid'

doc.add_heading(f'Returns', level=4)
doc.add_paragraph(f'<Paste the regional bar graph from Motiro report here>')
                                                                            
doc.add_heading(f'Rewards', level=4)
doc.add_paragraph(f'<Paste the regional bar graph from Motiro report here>')

doc.add_heading(f'Status', level=4)
doc.add_paragraph(f'<Paste the regional bar graph from Motiro report here>')

doc.add_heading(f'Key result', level=4)
doc.add_paragraph(f'Team discussion findings on management and extrinsic motivators')
doc.add_paragraph(f'Describe and explain patterns, commonalities and differences between teams in the region including relevant elements from their plans of action') 

doc.add_heading(f'Pathways to improved motivation in {structure_name}', level=3)

doc.add_paragraph('The figure below shows the most important correlations in the regional survey data. Because Motiro is based on SDT, the correlations are likely to be causal influences.')

doc.add_picture(f'{NS} {structure_name} SDTCorrNetworkGraph.png', width=Inches(6), height=Inches(4))

doc.add_heading(f'Discussion on {structure_name} survey and team discussion findings', level=3)
doc.add_paragraph(f'What are the main problems in motivation, engagement and retention in {structure_name} as a whole?')
doc.add_paragraph(f'What do the commonatities and differences between teams in {structure_name} inform us about how motivation, wellbeing, engagement and retention can be improved?')

doc.add_heading(f'Recommendations and next steps', level=2)
doc.add_heading(f'Recommendations from the teams', level=3)
doc.add_paragraph(f'What main solutions have the teams identified for the problems relevant for the region as a whole (especially solutions mentioned by several teams)? What support will they need to implement these solutions? For each recommendation, refer to a result or a pattern among teams in the report. If the recommendation is not borne out by the survey and discussion results, then it has no place here.')

doc.add_heading(f'Recommendations from the {structure}', level=3)
doc.add_paragraph(f'What do you recommend to the teams? What support will you be providing to the teams so that they can improve motivation? What support does the {structure_name} need to strengthen volunteer and staff motivation in your region?')

doc.add_heading(f'Next steps', level=3)
doc.add_paragraph(f'What are the next steps for the {structure_name} and the teams in the region? What will be the timeline for these steps?')

doc.add_heading('Conclusion', level=2)
doc.add_paragraph('we summarize the main findings of the report and the main recommendations for the teams and the {structure_name}.')

# Save Document
doc.save(f'{structure_name} report {month} {year}.docx')
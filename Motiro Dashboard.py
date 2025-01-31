import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_hyperlink(paragraph, text, url):
    """
    Add a hyperlink to a paragraph.
    :param paragraph: The paragraph to add the hyperlink to.
    :param text: The text to display for the hyperlink.
    :param url: The URL to link to.
    """
    # Create the hyperlink tag and add required attributes
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a run for the hyperlink text
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Apply formatting to the hyperlink
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    # Add underline formatting
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    run.append(rPr)

    # Add the text to the run
    t = OxmlElement('w:t')
    t.text = text
    run.append(t)
    hyperlink.append(run)

    # Append the hyperlink to the paragraph
    paragraph._p.append(hyperlink)

# Create a Word Document
doc = Document()

#Set Default Font and Size for the Document
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10)

#Add a custom style for headings and set the font properties
heading_style = doc.styles.add_style('CustomHeading', 1)
heading_font = heading_style.font
heading_font.name = 'Arial'
heading_font.size = Pt(12)
heading_font.bold = True
heading_font.color.rgb = RGBColor(0, 153, 153)  # CE Teal

#Add a custom style for the title and set the font properties
title_style = doc.styles.add_style('CustomTitle', 1)
title_font = title_style.font
title_font.name = 'Arial'
title_font.size = Pt(24)
title_font.bold = True
title_font.color.rgb = RGBColor(0, 111, 108)  # Dark blue

#Add a custom style for paragraphs and set the font properties
paragraph_style = doc.styles.add_style('CustomParagraph', 1)
paragraph_font = paragraph_style.font
paragraph_font.name = 'Arial'
paragraph_font.size = Pt(10)
paragraph_font.color.rgb = RGBColor(0, 0, 0)  # Black    


# Load the data
file_path = 'Individual.csv'
df = pd.read_csv(file_path, sep=",", encoding='utf-8')

## select which team, circle, category of respondent, etc. to display
#team='Volunteer'

# Filter the dataframe to keep only volunteers
#df = df[df['Position'] == team]
#df.info()

# Count the total number of respondents in filtered dataframe (n)
n = df['Volunteer'].sum() + df['Staff'].sum()

# create date variable from the 'Survey Data' column using the most recent time stamp.
df['Survey Data'] = pd.to_datetime(df['Survey Data'], format='mixed', errors='coerce')
most_recent_date = df['Survey Data'].max().strftime('%d-%B-%Y')

# Calculate statistics
total_volunteers = df['Volunteer'].sum()
total_staff = df['Staff'].sum()
# count unique occurence of teams and countries
total_teams = df['Team Name'].nunique()
total_countries = df['Country'].nunique()
# enter the app's current number of languages
total_languages = 11


# Add a title with custom formatting
title = doc.add_heading('Motiro Dashboard', level=1)
title.style = title_style

heading = doc.add_heading('Motiro app usage', level=2)
heading.style = heading_style

# Add a paragraph with a hyperlink
paragraph = doc.add_paragraph('The ')
add_hyperlink(paragraph, 'Motiro app', 'https://motiro.com')
paragraph.add_run(' works in multiple languages including the four official IFRC languages.')
#paragraph.style = paragraph_style

#doc.add_heading('Motiro Status Report', level=1)
#doc.add_heading('Motiro usage data', level=2)
# Add Text
doc.add_paragraph(f'As of {most_recent_date}, the Motiro app has been used by:')
paragraph.style = paragraph_style
doc.add_paragraph(f'{total_volunteers} volunteers', style='ListBullet')
paragraph.style = paragraph_style
doc.add_paragraph(f'{total_staff} staff', style='ListBullet')
paragraph.style = paragraph_style
doc.add_paragraph(f'from {total_teams} teams', style='ListBullet')
paragraph.style = paragraph_style
doc.add_paragraph(f'belonging to {total_countries} RCRC entities.', style='ListBullet')
paragraph.style = paragraph_style

# Add bar distribution plots by country and by region

# Add a table with one row and two columns
table = doc.add_table(rows=1, cols=2)

# Add the first picture to the first cell
cell1 = table.cell(0, 0)
paragraph1 = cell1.paragraphs[0]
run1 = paragraph1.add_run()
run1.add_picture('RespondentsByCountrySorted.png', width=Inches(3), height=Inches(3))

# Add the second picture to the second cell
cell2 = table.cell(0, 1)
paragraph2 = cell2.paragraphs[0]
run2 = paragraph2.add_run()
run2.add_picture('RespondentsByRegionSorted.png', width=Inches(3), height=Inches(3))

# Add time series Plots

# Add a table with one row and two columns
table = doc.add_table(rows=1, cols=2)

# Add the first picture to the first cell
cell1 = table.cell(0, 0)
paragraph1 = cell1.paragraphs[0]
run1 = paragraph1.add_run()
run1.add_picture('responses_over_time.png', width=Inches(3), height=Inches(2.5))

# Add the second picture to the second cell
cell2 = table.cell(0, 1)
paragraph2 = cell2.paragraphs[0]
run2 = paragraph2.add_run()
run2.add_picture('responses_over_time_cumulative.png', width=Inches(3), height=Inches(2.5))

doc.add_page_break()
heading = doc.add_heading('Motiro findings', level=2)
heading.style = heading_style

doc.add_paragraph('What is the quality of volunteer and staff motivation and engagement?')
paragraph.style = paragraph_style

# Add a table with one row and two columns
table = doc.add_table(rows=1, cols=2)

# Add the first picture to the first cell
cell1 = table.cell(0, 0)
paragraph1 = cell1.paragraphs[0]
run1 = paragraph1.add_run()
run1.add_picture('Volunteer_spider.png', width=Inches(3), height=Inches(3))

# Add the second picture to the second cell
cell2 = table.cell(0, 1)
paragraph2 = cell2.paragraphs[0]
run2 = paragraph2.add_run()
run2.add_picture('Staff_spider.png', width=Inches(3), height=Inches(3))

# Add a ligne break
doc.add_paragraph('')
paragraph.style = paragraph_style

doc.add_paragraph('What are the pathways toward improved motivation, engagement and well-being?')
paragraph.style = paragraph_style
# Add a table with one row and two columns
table = doc.add_table(rows=1, cols=2)

# Add the first picture to the first cell
cell1 = table.cell(0, 0)
paragraph1 = cell1.paragraphs[0]
run1 = paragraph1.add_run()
run1.add_picture('Volunteer SDTCorrNetworkGraph.png', width=Inches(3), height=Inches(2))

# Add the second picture to the second cell
cell2 = table.cell(0, 1)
paragraph2 = cell2.paragraphs[0]
run2 = paragraph2.add_run()
run2.add_picture('Staff SDTCorrNetworkGraph.png', width=Inches(3), height=Inches(2))

# Save Document
doc.save('MotiroDashboard.docx')
